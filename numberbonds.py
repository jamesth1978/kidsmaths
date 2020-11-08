import random
import argparse
import collections
from docx import Document
from docx.shared import Pt
from docx.shared import Mm
import os

doc = Document()
section = doc.sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)
section.left_margin = Mm(25.4)
section.right_margin = Mm(25.4)
section.top_margin = Mm(25.4)
section.bottom_margin = Mm(25.4)
section.header_distance = Mm(12.7)
section.footer_distance = Mm(12.7)
fontsize = 13

p = doc.add_paragraph()
run = p.add_run()
font = run.font
font.size = Pt(fontsize)


parser = argparse.ArgumentParser(description='Print some number bond questions')
parser.add_argument('target', type=int, help='The sum you want to target')
parser.add_argument('num', type=int, help='The number of questions')
parser.add_argument('double_spaced', type=str, help='Print double spaced lines (y/n)')

args = parser.parse_args()

target = args.target
num_questions = args.num
double_spaced = args.double_spaced

lines = []

previous = 999
i = 0
recents = collections.deque([],3)


while i <= num_questions:
    number = random.randint(0,target)
        
    if number not in recents: 
        lines.append(str(number) + " + " + " _____________  = " + str(target))
        i+=1
    
    recents.append(number)
    

        
if double_spaced == 'y':
    linebreak = "\r\n\r\n"
else:
    linebreak = "\r\n"
    
j = 0
while j <= num_questions-1:
    if j==(num_questions-2):
        line_end = ""
    else:
        line_end = linebreak
    combo_line = lines[j] + "\t \t \t \t \t" + lines[j+1]+line_end
    print(combo_line)
    run.text = combo_line
    run = p.add_run()
    font = run.font
    font.size = Pt(fontsize)
    j+=2

doc.save('numberbonds.docx')
os.system('unoconv --stdout numberbonds.docx |lpr -U pi -P MP280-series')
    
    

