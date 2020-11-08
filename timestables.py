import random
import sys
import argparse
import collections
from docx import Document
from docx.shared import Pt
from docx.shared import Mm
import os

doc = Document()
section = doc.sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)
section.left_margin = Mm(25.4)
section.right_margin = Mm(25.4)
section.top_margin = Mm(25.4)
section.bottom_margin = Mm(25.4)
section.header_distance = Mm(12.7)
section.footer_distance = Mm(12.7)
fontsize = 13

p = doc.add_paragraph()
run = p.add_run()
font = run.font
font.size = Pt(fontsize)


parser = argparse.ArgumentParser(description='Print some times tables questions')
parser.add_argument('factor', type=int, help='The times table you are after, usually between 1 and 12')
parser.add_argument('up_to', type=int, help='The "up to" limit, usually 12')
parser.add_argument('num', type=int, help='The number of questions')
parser.add_argument('mix', type=str, help='Mix of factors (y/n)')
if len(sys.argv)==1:
    parser.print_help()
    parser.exit()
args = parser.parse_args()

limit = args.up_to
factor = args.factor
num_questions = args.num
mixture = args.mix

if limit <3:
    print("The \"up-to\" limit must be greater than 2.  Exiting.")
    raise SystemExit

lines = []

i = 0
recentsl = collections.deque([],3)
recentsr = collections.deque([],3)
increment = 0

if mixture == 'y':
    while i <= num_questions:
        left = random.randint(0,limit)
        right = random.randint(0,limit)
        if left not in recentsl:
            l_increment = 1
        if right not in recentsr:
            r_increment = 1

        if l_increment & r_increment == 1:
            i += 1
            lines.append(str(left) + " x " + str(right) + " = ")
            recentsl.append(left)
            recentsr.append(right)
                
if mixture == 'n':
    i = 0
    recents = collections.deque([],3)
    right = factor
      
    while i <= num_questions:
        left = random.randint(0,limit)
        if left not in recents:
            lines.append(str(left) + " x " + str(right) + " = ")
            i += 1
            recents.append(left)

       
        
j = 0
while j <= num_questions-1:
    if(j==(num_questions-2)):
        line_end = ""
    else:
        line_end = "\r\n"
    combo_line = lines[j] + "\t \t \t \t \t \t" + lines[j+1]+line_end
    print(combo_line)
    run.text = combo_line
    run = p.add_run()
    font = run.font
    font.size = Pt(fontsize)
    j+=2

doc.save('timestable.docx')
os.system('unoconv --stdout timestable.docx |lpr -U pi -P MP280-series')   
    

